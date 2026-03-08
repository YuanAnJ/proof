from docx import Document
from docx.shared import Pt,RGBColor,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from forms.utils.afp import add_float_pic
from forms.models import FormsData

def generate_doc(data):
    feedback_date = data.feedback_date.strftime('%Y年%m月')
    feedback_department = data.feedback_department
    title = data.title
    authors_info = FormsData.get_author_info(data)
    accept_level = data.accept_level
    instruction_level = data.instruction_level
    remark = data.remark

    timestamp = datetime.now().strftime("%Y年%m月%d日")

    # 创建一个新的文档对象
    doc = Document()

    # Set paper size to A4 (21.0 x 29.7 cm)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.60)
    section.bottom_margin = Cm(2.05)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 添加段落
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置段落居中
    para.paragraph_format.space_before = Pt(0)  # 设置段前间距为0
    para.paragraph_format.space_after = Pt(0)   # 设置段后间距为0
    para.paragraph_format.line_spacing = Pt(40)  # 设置行距为固定值40磅
    run = para.add_run('深      圳      大      学')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "华文中宋")  # 中文字体
    run.font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色
    run.font.size = Pt(42)  # 初号对应42磅

    # 添加图片
    picture1 = doc.add_picture('media/line.png')  # 如果有图片，取消注释并提供图片路径
    picture1.width = Cm(15.5)  # 设置图片宽度为15.5厘米

    # 添加段落2
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置段落居左
    para2.paragraph_format.space_before = Pt(0)  # 设置段前间距为0
    para2.paragraph_format.space_after = Pt(0)   # 设置段后间距为0
    run2 = para2.add_run(f'内部文件                                                           编号：{data.number}')
    run2.font.name = 'Times New Roman'
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")  # 中文字体
    run2.font.size = Pt(16)  # 3号字体对应16磅

    # 添加段落3
    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置段落居中
    para3.paragraph_format.space_before = Pt(0)  # 设置段前间距为0
    para3.paragraph_format.space_after = Pt(0)   # 设置段后间距为0
    run3 = para3.add_run('\n信息采用证明')
    run3.font.name = 'Times New Roman'
    run3._element.rPr.rFonts.set(qn("w:eastAsia"), "方正小标宋简体")  # 中文字体
    run3.font.size = Pt(22)  # 2号字体对应22磅

    # 添加段落4
    para4 = doc.add_paragraph()
    para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
    para4.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅
    
    # 生成作者和单位文本
    text = f'根据上级党政部门{feedback_date}信息情况反馈，兹证明'
    
    # 根据authors_info生成作者列表文本
    authors_text = []
    for info in authors_info:
        authors_text.append(f"{info['unit']}{info['author']}")
    
    # 用顿号连接作者信息
    if authors_text:
        text += '、'.join(authors_text)
        text += f'撰写的咨询报告《{title}》'
    
    # 根据批示级别和采纳级别生成不同的文本
    accept = False
    instruction = False
    if accept_level=='无' or accept_level is None:
        accept = True
    if instruction_level=='无' or instruction_level is None:
        instruction = True
    if accept_level != '无' and instruction_level != '无' and accept_level is not None and instruction_level is not None:
        text += f'获{accept_level}内参采用，并获{instruction_level}肯定性批示，为有关部门和领导决策提供积极参考。'
    elif accept and instruction_level != '无':
        text += f'获{instruction_level}肯定性批示，为有关部门和领导决策提供积极参考。'
    elif instruction and accept_level != '无':
        text += f'获{accept_level}内参采用，为有关部门和领导决策提供积极参考。'
    # 添加文本
    run4 = para4.add_run(text)
    run4.font.name = '仿宋'
    run4._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 中文字体
    run4.font.size = Pt(16)  # 3号字体对应16磅

    # 添加段落5
    para5 = doc.add_paragraph()
    para5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
    para5.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅
    run5 = para5.add_run('此证明仅供职称评审、课题结项、评奖评优、绩效考核使用，')
    run5.font.size = Pt(16)  # 3号字体对应16磅
    run5.font.name = '仿宋'
    run5._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 中文字体
    run6 = para5.add_run('请妥善保管，勿拍照、公开发布、宣传或网上传播等，按规定做好保密工作。')
    run6.font.size = Pt(16)  # 3号字体对应16磅
    run6.font.name = '仿宋'
    run6._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 中文字体
    run6.bold = True

    # 添加段落6
    para6 = doc.add_paragraph()
    para6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 设置段落居右
    para6.paragraph_format.first_line_indent = Pt(32)  # 首行缩进32磅
    run7 = para6.add_run('特此证明。')
    run7.font.name = '仿宋'
    run7._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 中文字体
    run7.font.size = Pt(16)  # 3号字体对应16磅

    # 添加段落7
    para7 = doc.add_paragraph()
    para7.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 设置段落居右
    run8 = para7.add_run(f'\n\n深圳大学社会科学部\n{timestamp}')
    run8.font.name = '仿宋'
    run8._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 中文字体
    run8.font.size = Pt(16)  # 3号字体对应16磅
    
    # 添加公章
    add_float_pic(para7, 'media/cachet.png', width=Cm(4.0), height=Cm(4.0), pos_x=Cm(14.18), pos_y=Cm(17.34))

    # 保存文档
    file_name = f'{data.number}'
    doc.save('media/proof_file.docx')

    return file_name