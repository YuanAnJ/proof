import re
from datetime import datetime

from docx import Document


def _clean_text(value):
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\u3000", " ")
    return re.sub(r"\s+", " ", text).strip()


def _collect_doc_text(doc):
    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)

    return "\n".join(lines)


def _extract_date_from_context(text):
    """从'根据上级党政部门2017年02月信息情况反馈'格式中提取日期"""
    pattern = r"根据上级党政部门\s*(\d{4})\s*年\s*(\d{1,2})\s*月"
    match = re.search(pattern, text)
    if match:
        year = match.group(1)
        month = match.group(2).zfill(2)
        return {
            "feedback_date": "",
            "feedback_date_raw": f"{year}-{month}",
            "is_partial_date": True,
        }
    return {
        "feedback_date": "",
        "feedback_date_raw": "",
        "is_partial_date": False,
    }


def _extract_title_from_context(text):
    """从'报告《有关于咨询决策的报告》'格式中提取题目"""
    pattern = r"报告[《](.*?)[》]"
    match = re.search(pattern, text)
    if match:
        return _clean_text(match.group(1))
    
    # 兼容其他书名号格式
    pattern2 = r"[《](.*?)[》]"
    match2 = re.search(pattern2, text)
    if match2:
        return _clean_text(match2.group(1))
    
    return ""


def _extract_authors_from_context(text):
    """从'深圳大学小纪、厦门大学小李、复旦大学小赵'格式中提取作者和单位
    返回格式：[(姓名, 单位), ...]
    """
    authors = []
    
    # 查找包含作者信息的段落（通常在"兹证明"之后）
    author_pattern = r"兹证明(.+?)撰写"
    match = re.search(author_pattern, text, re.DOTALL)
    if not match:
        return authors
    
    author_text = match.group(1)
    
    # 解析"单位+姓名"格式，用顿号分隔
    # 匹配模式：XXX大学/XXX单位 + 姓名
    pieces = re.split(r"[、，]", author_text)
    
    for piece in pieces:
        piece = _clean_text(piece)
        if not piece:
            continue
        
        # 尝试匹配"单位名+人名"格式
        # 假设人名通常是最后2-4个字符
        if len(piece) > 2:
            # 查找常见单位后缀
            unit_match = re.match(r"(.+?(?:大学|学院|研究院|研究所|中心|部|局|委|厅|司|处|科|公司|集团))(.+)", piece)
            if unit_match:
                unit = _clean_text(unit_match.group(1))
                name = _clean_text(unit_match.group(2))
                if unit and name:
                    authors.append((name, unit))
                    continue
            
            # 如果没有明确单位标识，尝试从后往前取2-4个字符作为姓名
            for name_len in [2, 3, 4]:
                if len(piece) > name_len:
                    name = piece[-name_len:]
                    unit = piece[:-name_len]
                    # 验证单位名不为空且有合理长度
                    if len(unit) >= 2:
                        authors.append((name, unit))
                        break
    
    return authors


def _extract_accept_level_from_context(text):
    """从'获中办内参采用'格式中提取采纳级别"""
    pattern = r"获(.+?)内参采用"
    match = re.search(pattern, text)
    if match:
        level = _clean_text(match.group(1))
        # 移除可能的"、"分隔符
        levels = [l.strip() for l in re.split(r"[、，,]", level) if l.strip()]
        return "、".join(levels)
    return ""


def _extract_instruction_level_from_context(text):
    """从'获中央主要领导肯定性批示'格式中提取批示级别"""
    # 使用非贪婪匹配，并且要求"获"后面紧跟的不是"内参采用"
    pattern = r"[并、，获]\s*([^获、，内]+?)肯定性批示"
    match = re.search(pattern, text)
    if match:
        level = _clean_text(match.group(1))
        # 移除可能的"、"分隔符和"获"字
        level = level.replace("获", "").strip()
        levels = [l.strip() for l in re.split(r"[、，,]", level) if l.strip()]
        return "、".join(levels)
    return ""


def _parse_date_value(raw):
    """兼容直接提供的日期值"""
    value = _clean_text(raw)
    if not value:
        return {
            "feedback_date": "",
            "feedback_date_raw": "",
            "is_partial_date": False,
        }

    normalized = value.replace("年", "-").replace("月", "-").replace("日", "")
    normalized = normalized.replace("/", "-").strip("-")

    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
        try:
            dt = datetime.strptime(normalized, fmt)
            return {
                "feedback_date": dt.strftime("%Y-%m-%d"),
                "feedback_date_raw": value,
                "is_partial_date": False,
            }
        except ValueError:
            pass

    if re.match(r"^\d{4}-\d{1,2}$", normalized):
        return {
            "feedback_date": "",
            "feedback_date_raw": normalized,
            "is_partial_date": True,
        }

    return {
        "feedback_date": "",
        "feedback_date_raw": value,
        "is_partial_date": False,
    }


def parse_docx_file(file_obj):
    doc = Document(file_obj)
    full_text = _collect_doc_text(doc)

    record = {
        "feedback_date": "",
        "feedback_date_raw": "",
        "is_partial_date": False,
        "feedback_department": "",  # 默认留空
        "title": "",
        "author": "",
        "unit": "",
        "accept_level": "",
        "instruction_level": "",
        "remark": "",  # 默认留空
        "other_accept_level": "",
        "other_instruction_level": "",
    }

    # 1. 从正文中提取反馈日期（格式：根据上级党政部门2017年02月信息情况反馈）
    date_info = _extract_date_from_context(full_text)
    record.update(date_info)

    # 2. 反馈部门默认留空（根据用户说明）

    # 3. 从正文中提取题目（格式：报告《XXX》）
    record["title"] = _extract_title_from_context(full_text)

    # 4. 从正文中提取作者和单位（格式：深圳大学小纪、厦门大学小李）
    authors = _extract_authors_from_context(full_text)
    
    if authors:
        first_author, first_unit = authors[0]
        record["author"] = first_author
        record["unit"] = first_unit

    for idx in range(2, 11):
        record[f"author_{idx}"] = ""
        record[f"unit_{idx}"] = ""

    for idx, pair in enumerate(authors[1:10], start=2):
        record[f"author_{idx}"] = pair[0]  # name
        record[f"unit_{idx}"] = pair[1]    # unit

    # 5. 从正文中提取采纳级别（格式：获中办内参采用）
    record["accept_level"] = _extract_accept_level_from_context(full_text)

    # 6. 从正文中提取批示级别（格式：获中央主要领导肯定性批示）
    record["instruction_level"] = _extract_instruction_level_from_context(full_text)

    # 7. 备注默认留空（根据用户说明）

    return record
